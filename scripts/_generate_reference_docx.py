#!/usr/bin/env python3
"""_generate_reference_docx.py — build assets/styles/lawyer-reference.docx.

Pandoc uses this file's styles when rendering markdown to Word
(`pandoc --reference-doc=...`). Re-run this when the Word styling needs to change.

Style spec:
- Body: Calibri 11
- Headings: Calibri Light, sized 18 / 14 / 12 / 11, navy #1F3A5F
- Margins: 2.5 cm
- Page numbers: bottom-centre
- List indents, blockquote indent, table grid

This file is checked into the repo so end users do not need to run the
generator themselves.
"""
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor
from lxml import etree

OUT = Path(__file__).resolve().parent.parent / "assets" / "styles" / "lawyer-reference.docx"

NAVY = RGBColor(0x1F, 0x3A, 0x5F)


def set_paragraph_style(style, *, font="Calibri", size=11, bold=False, color=None,
                        space_before_pt=0, space_after_pt=6, keep_with_next=False):
    f = style.font
    f.name = font
    f.size = Pt(size)
    f.bold = bold
    if color is not None:
        f.color.rgb = color
    pf = style.paragraph_format
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    if keep_with_next:
        pf.keep_with_next = True


def add_page_number_footer(section):
    """Drop a centred PAGE field into the section footer."""
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Clear any existing runs
    for r in list(p.runs):
        r.text = ""
    # Build the PAGE field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run = p.add_run()
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def set_table_grid(doc):
    """Make sure the default table style has visible grid lines."""
    try:
        _ = doc.styles["Table Grid"]
    except KeyError:
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        body = doc.element.body
        body.remove(tbl._element)


# --------------------------------------------------------------------------
# Pandoc's named styles must exist in the reference doc, otherwise pandoc
# falls back to defaults that ignore our spec.
# --------------------------------------------------------------------------
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _w(tag):
    return f"{{{W}}}{tag}"


def _w_el(tag, **attrs):
    el = OxmlElement(f"w:{tag}")
    for k, v in attrs.items():
        el.set(_w(k), v)
    return el


def add_verbatim_char_style(doc):
    """Add the 'VerbatimChar' character style pandoc uses for inline `code`."""
    styles_el = doc.styles.element
    # Skip if already present
    for s in styles_el.findall(_w("style")):
        if s.get(_w("styleId")) == "VerbatimChar":
            return

    style = OxmlElement("w:style")
    style.set(_w("type"), "character")
    style.set(_w("customStyle"), "1")
    style.set(_w("styleId"), "VerbatimChar")

    name = OxmlElement("w:name")
    name.set(_w("val"), "Verbatim Char")
    style.append(name)

    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(_w("ascii"), "Consolas")
    rFonts.set(_w("hAnsi"), "Consolas")
    rFonts.set(_w("cs"), "Consolas")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(_w("val"), "20")  # 10pt = 20 half-points
    rPr.append(sz)
    style.append(rPr)
    styles_el.append(style)


def add_source_code_paragraph_style(doc):
    """Add the 'Source Code' paragraph style pandoc uses for fenced code blocks."""
    styles_el = doc.styles.element
    for s in styles_el.findall(_w("style")):
        if s.get(_w("styleId")) == "SourceCode":
            return

    style = OxmlElement("w:style")
    style.set(_w("type"), "paragraph")
    style.set(_w("customStyle"), "1")
    style.set(_w("styleId"), "SourceCode")

    name = OxmlElement("w:name")
    name.set(_w("val"), "Source Code")
    style.append(name)

    pPr = OxmlElement("w:pPr")
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(_w("val"), "single")
        b.set(_w("sz"), "4")
        b.set(_w("space"), "1")
        b.set(_w("color"), "DDDDDD")
        pBdr.append(b)
    pPr.append(pBdr)
    shd = OxmlElement("w:shd")
    shd.set(_w("val"), "clear")
    shd.set(_w("color"), "auto")
    shd.set(_w("fill"), "F5F5F5")
    pPr.append(shd)
    style.append(pPr)

    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(_w("ascii"), "Consolas")
    rFonts.set(_w("hAnsi"), "Consolas")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(_w("val"), "20")
    rPr.append(sz)
    style.append(rPr)
    styles_el.append(style)


def add_pandoc_table_style(doc):
    """Add the 'Table' table style pandoc references with <w:tblStyle w:val="Table"/>.

    Bold first row, grid borders, light alternating row fill.
    """
    styles_el = doc.styles.element
    for s in styles_el.findall(_w("style")):
        if s.get(_w("styleId")) == "Table":
            return

    style = OxmlElement("w:style")
    style.set(_w("type"), "table")
    style.set(_w("customStyle"), "1")
    style.set(_w("styleId"), "Table")

    name = OxmlElement("w:name")
    name.set(_w("val"), "Table")
    style.append(name)
    base = OxmlElement("w:basedOn")
    base.set(_w("val"), "TableNormal")
    style.append(base)

    # rPr for cell text
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(_w("ascii"), "Calibri")
    rFonts.set(_w("hAnsi"), "Calibri")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(_w("val"), "22")  # 11pt
    rPr.append(sz)
    style.append(rPr)

    # Table-wide properties: grid borders
    tblPr = OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(_w("val"), "single")
        b.set(_w("sz"), "4")
        b.set(_w("space"), "0")
        b.set(_w("color"), "BFBFBF")
        tblBorders.append(b)
    tblPr.append(tblBorders)
    tblCellMar = OxmlElement("w:tblCellMar")
    for side, val in (("top", "60"), ("left", "100"), ("bottom", "60"), ("right", "100")):
        m = OxmlElement(f"w:{side}")
        m.set(_w("w"), val)
        m.set(_w("type"), "dxa")
        tblCellMar.append(m)
    tblPr.append(tblCellMar)
    style.append(tblPr)

    # First-row conditional formatting: bold + navy fill
    tblStylePr = OxmlElement("w:tblStylePr")
    tblStylePr.set(_w("type"), "firstRow")
    rPr2 = OxmlElement("w:rPr")
    bold = OxmlElement("w:b"); bold.set(_w("val"), "1")
    rPr2.append(bold)
    color = OxmlElement("w:color"); color.set(_w("val"), "FFFFFF")
    rPr2.append(color)
    tblStylePr.append(rPr2)
    tcPr = OxmlElement("w:tcPr")
    shd = OxmlElement("w:shd")
    shd.set(_w("val"), "clear")
    shd.set(_w("color"), "auto")
    shd.set(_w("fill"), "1F3A5F")
    tcPr.append(shd)
    tblStylePr.append(tcPr)
    style.append(tblStylePr)

    # Banded rows: light grey on even rows
    bandStyle = OxmlElement("w:tblStylePr")
    bandStyle.set(_w("type"), "band2Horz")
    band_tcPr = OxmlElement("w:tcPr")
    band_shd = OxmlElement("w:shd")
    band_shd.set(_w("val"), "clear")
    band_shd.set(_w("color"), "auto")
    band_shd.set(_w("fill"), "F5F5F5")
    band_tcPr.append(band_shd)
    bandStyle.append(band_tcPr)
    style.append(bandStyle)

    styles_el.append(style)


def add_compact_style(doc):
    """Pandoc uses 'Compact' paragraph style inside list items / tight tables."""
    styles_el = doc.styles.element
    for s in styles_el.findall(_w("style")):
        if s.get(_w("styleId")) == "Compact":
            return

    style = OxmlElement("w:style")
    style.set(_w("type"), "paragraph")
    style.set(_w("customStyle"), "1")
    style.set(_w("styleId"), "Compact")
    name = OxmlElement("w:name")
    name.set(_w("val"), "Compact")
    style.append(name)
    base = OxmlElement("w:basedOn")
    base.set(_w("val"), "Normal")
    style.append(base)

    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(_w("after"), "0")
    spacing.set(_w("before"), "0")
    pPr.append(spacing)
    style.append(pPr)
    styles_el.append(style)


def build():
    doc = Document()

    # Section / page setup
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        add_page_number_footer(section)

    # Normal / body
    normal = doc.styles["Normal"]
    set_paragraph_style(normal, font="Calibri", size=11, space_after_pt=6)

    # Headings
    set_paragraph_style(doc.styles["Heading 1"], font="Calibri Light", size=18,
                        bold=True, color=NAVY, space_before_pt=18, space_after_pt=8,
                        keep_with_next=True)
    set_paragraph_style(doc.styles["Heading 2"], font="Calibri Light", size=14,
                        bold=True, color=NAVY, space_before_pt=12, space_after_pt=6,
                        keep_with_next=True)
    set_paragraph_style(doc.styles["Heading 3"], font="Calibri Light", size=12,
                        bold=False, color=NAVY, space_before_pt=10, space_after_pt=4,
                        keep_with_next=True)
    set_paragraph_style(doc.styles["Heading 4"], font="Calibri Light", size=11,
                        bold=False, color=NAVY, space_before_pt=8, space_after_pt=4,
                        keep_with_next=True)

    # Title (used by pandoc when --metadata title=...)
    try:
        title_style = doc.styles["Title"]
    except KeyError:
        title_style = None
    if title_style is not None:
        set_paragraph_style(title_style, font="Calibri Light", size=24,
                            bold=True, color=NAVY, space_after_pt=12)

    # Lists — pandoc uses "List Bullet" / "List Number"
    for sname in ["List Bullet", "List Number", "List Paragraph"]:
        try:
            s = doc.styles[sname]
            set_paragraph_style(s, font="Calibri", size=11, space_after_pt=4)
            s.paragraph_format.left_indent = Cm(0.6)
        except KeyError:
            pass

    # Block quote — pandoc uses "Block Text" or "Quote"
    for sname in ["Quote", "Block Text", "Intense Quote"]:
        try:
            s = doc.styles[sname]
            set_paragraph_style(s, font="Calibri", size=11, space_after_pt=6)
            s.paragraph_format.left_indent = Cm(1.0)
            s.paragraph_format.right_indent = Cm(0.5)
            # italic to set emphasis on quoted regulation text
            s.font.italic = True
        except KeyError:
            pass

    # Code styles
    for sname in ["Source Code", "Verbatim Char", "Code"]:
        try:
            s = doc.styles[sname]
            s.font.name = "Consolas"
            s.font.size = Pt(10)
        except KeyError:
            pass

    # Make sure Table Grid is registered
    set_table_grid(doc)

    # Pandoc-specific styles
    add_verbatim_char_style(doc)
    add_source_code_paragraph_style(doc)
    add_pandoc_table_style(doc)
    add_compact_style(doc)

    # Add a single placeholder paragraph so pandoc has something to attach to.
    # Pandoc replaces document body content with the rendered markdown; only
    # styles + section properties + headers/footers from the reference doc
    # are reused.
    p = doc.add_paragraph("Reference template — replaced by pandoc on render.",
                          style="Normal")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
